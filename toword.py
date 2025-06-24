import docx
from docx.shared import RGBColor, Pt
from tqdm import tqdm
import os

def parse_mcqs_from_txt(filepath):
    """
    Parses a structured text file to extract MCQs.
    
    Args:
        filepath (str): The path to the input .txt file.

    Returns:
        list: A list of dictionaries, where each dictionary represents one MCQ.
    """
    if not os.path.exists(filepath):
        print(f"Error: The file '{filepath}' was not found.")
        print("Please make sure the .txt file is in the same directory as the script.")
        return None

    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    mcqs = []
    current_mcq = {}

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("Question:"):
            # If we are already building an MCQ, save it before starting a new one
            if current_mcq:
                mcqs.append(current_mcq)
            current_mcq = {"question": line[10:].strip(), "options": []}
        elif line.startswith("-"):
            if "question" in current_mcq:
                current_mcq["options"].append(line[1:].strip())
        elif line.startswith("---"):
            # End of an MCQ block
            if current_mcq:
                mcqs.append(current_mcq)
            current_mcq = {}
            
    # Add the last MCQ if the file doesn't end with a separator
    if current_mcq and "question" in current_mcq:
        mcqs.append(current_mcq)
        
    return mcqs

def create_word_document(mcqs, output_filename):
    """
    Creates a formatted Word document from a list of MCQs.
    
    Args:
        mcqs (list): The list of MCQ dictionaries.
        output_filename (str): The name for the output .docx file.
    """
    doc = docx.Document()
    doc.add_heading("PakMCQs Complete Question Bank", level=0)
    
    p = doc.add_paragraph()
    p.add_run(f"A formatted collection of {len(mcqs)} questions.").italic = True
    doc.add_paragraph() # Spacer

    for i, mcq in enumerate(tqdm(mcqs, desc="Creating Word Document")):
        # Add the question (e.g., "1. Who is the PM?")
        p_question = doc.add_paragraph()
        p_question.add_run(f"{i + 1}. {mcq['question']}").bold = True
        
        # Add the options
        for option_text in mcq['options']:
            is_correct = "[Correct]" in option_text
            clean_option = option_text.replace("[Correct]", "").strip()

            p_option = doc.add_paragraph(clean_option, style='List Bullet')
            
            # If it's the correct answer, format the entire line
            if is_correct:
                run = p_option.runs[0]
                run.bold = True
                font = run.font
                font.color.rgb = RGBColor(0x00, 0x80, 0x00) # Green color

        doc.add_paragraph() # Add a space after each MCQ block

    doc.save(output_filename)


# --- Main Execution ---
INPUT_FILENAME = "pakmcqs_all_questions_fast.txt"
OUTPUT_FILENAME = "PakMCQs_Formatted_from_TXT.docx"

print(f"Reading MCQs from '{INPUT_FILENAME}'...")
mcq_data = parse_mcqs_from_txt(INPUT_FILENAME)

if mcq_data:
    print(f"Successfully parsed {len(mcq_data)} MCQs.")
    create_word_document(mcq_data, OUTPUT_FILENAME)
    print(f"\nProcess finished successfully!")
    print(f"Your formatted Word file is saved as: {OUTPUT_FILENAME}")